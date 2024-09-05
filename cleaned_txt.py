#using https://github.com/Dself76/prediction_pipline repo in github 
import nltk
import ssl
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
import os
import mysql.connector

# SSL workaround for NLTK downloads
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

# Download necessary NLTK data
nltk.download('wordnet')
nltk.download('omw-1.4')
nltk.download('punkt')
nltk.download('stopwords')

class CryptoDocumentProcessor:
    def __init__(self, db_config, file_paths):
        self.db_config = db_config
        self.file_paths = file_paths
        self.crypto_keywords = self.get_crypto_keywords()

    def get_crypto_keywords(self):
        conn = mysql.connector.connect(**self.db_config)
        cursor = conn.cursor()
        cursor.execute("SELECT name, ticker FROM CRYPTO_NAMES")
        crypto_keywords = {name.lower() for name, ticker in cursor.fetchall()}
        conn.close()
        return crypto_keywords

    def read_docx(self, file_path):
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return ' '.join(full_text)

    def clean_and_tokenize(self, text):
        try:
            text = text.lower()
            tokens = word_tokenize(text)
            stop_words = set(stopwords.words('english'))
            filtered_tokens = [word for word in tokens if word not in stop_words and word.isalnum()]
            return filtered_tokens
        except LookupError as e:
            print(f"Error: {e}")
            return []

    def process_documents(self):
        all_tokens = []
        for file_path in self.file_paths:
            try:
                if not os.path.exists(file_path):
                    print(f"File not found: {file_path}")
                    continue
                text = self.read_docx(file_path)
                tokens = self.clean_and_tokenize(text)
                all_tokens.append(tokens)
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")
        return all_tokens

    def find_crypto_keywords(self):
        all_tokens = self.process_documents()
        for i, tokens in enumerate(all_tokens):
            print(f"Tokens from document {i+1}:")
            if tokens:
                found_cryptos = set(tokens) & self.crypto_keywords
                if found_cryptos:
                    print(f"Found crypto keywords: {found_cryptos}")
                else:
                    print("No crypto keywords found")
            else:
                print("No tokens found")
            print("\n")

# Configuration for database connection
db_config = {
    'host': '127.0.0.1',
    'user': 'root',
    'password': 'Max123max123',
    'database': 'crypto_history'
}

# List of file paths
file_paths = ['data1.docx', 'BitcoindocforNLP.docx', 'data101.docx']

# Create an instance of the processor and find crypto keywords
processor = CryptoDocumentProcessor(db_config, file_paths)
processor.find_crypto_keywords()


'''
#these are the functions i created before switching to class

import nltk
import ssl
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
import os
import mysql.connector  # Import MySQL connector

# SSL workaround for NLTK downloads
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

# Download necessary NLTK data
nltk.download('wordnet')
nltk.download('omw-1.4')
nltk.download('punkt')
nltk.download('stopwords')

# Function to extract text from a .docx file
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return ' '.join(full_text)

# Function to clean and tokenize text
def clean_and_tokenize(text):
    try:
        text = text.lower()
        tokens = word_tokenize(text)
        stop_words = set(stopwords.words('english'))
        filtered_tokens = [word for word in tokens if word not in stop_words and word.isalnum()]
        return filtered_tokens
    except LookupError as e:
        print(f"Error: {e}")
        return []

# Connect to your database and retrieve crypto names and tickers
def get_crypto_keywords():
    conn = mysql.connector.connect(
        host='your_host',
        user='your_user',
        password='your_password',
        database='crypto_history'
    )
    cursor = conn.cursor()
    cursor.execute("SELECT name, ticker FROM CRYPTO_NAMES")
    crypto_keywords = {name.lower() for name, ticker in cursor.fetchall()}
    conn.close()
    return crypto_keywords

# List of file paths
file_paths = ['data1.docx', 'BitcoindocforNLP.docx', 'data101.docx']

# Retrieve crypto keywords from the database
crypto_keywords = get_crypto_keywords()

# Loop through each file path, read and process the document
all_tokens = []
for file_path in file_paths:
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
        text = read_docx(file_path)
        tokens = clean_and_tokenize(text)
        all_tokens.append(tokens)
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")

# Check for crypto keywords in the tokens
for i, tokens in enumerate(all_tokens):
    print(f"Tokens from document {i+1}:")
    if tokens:
        found_cryptos = set(tokens) & crypto_keywords
        if found_cryptos:
            print(f"Found crypto keywords: {found_cryptos}")
        else:
            print("No crypto keywords found")
    else:
        print("No tokens found")
    print("\n")
'''